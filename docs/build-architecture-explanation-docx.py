from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "DECEPTR_Explication_Architectures_Etapes_Objectifs.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 98, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E2F0D9"
LIGHT_ORANGE = "FCE4D6"
LIGHT_PURPLE = "EDE7F6"
LIGHT_GRAY = "F2F4F7"


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("DECEPTR v1 MVP - Explication des architectures")
    set_font(r, size=8.5, color=MUTED)


def para(doc, text="", bold=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, bold=bold, color=color, size=size)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r, size=10.5)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r, size=10.5)


def table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        set_font(r, bold=True, color=NAVY, size=9.2)
        shade(hdr[i], header_fill)
        margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            set_font(r, size=8.8)
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return t


def callout(doc, title, text, fill=LIGHT_GRAY):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade(cell, fill)
    margins(cell, top=130, bottom=130)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    set_font(r, bold=True, color=NAVY, size=10.2)
    r = p.add_run(text)
    set_font(r, size=9.5)
    doc.add_paragraph()


def cover(doc):
    para(doc, "CHAMBRE DES REPRESENTANTS DU MAROC", bold=True, color=BLUE, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "DECEPTR v1 MVP", bold=True, color=NAVY, size=26, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Explication simple des architectures, des etapes et des objectifs", color=MUTED, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    table(
        doc,
        ["Champ", "Valeur"],
        [
            ["Sujet", "Architecture reseau, architecture fonctionnelle et architecture technique"],
            ["Objectif", "Expliquer chaque composant, chaque flux et le choix de conception"],
            ["Projet", "Plateforme de cyberdeception et honeypots DECEPTR"],
            ["Contexte", "PFA - Chambre des Representants du Maroc"],
            ["Version", "v1.0 finale - Juin 2026"],
        ],
        widths=[1.5, 5.4],
    )
    callout(
        doc,
        "But du document",
        "Ce document est fait pour expliquer le schema d'architecture de maniere simple. Il aide a presenter le projet devant un jury ou un encadrant: quoi fait chaque outil, pourquoi il est la, et pourquoi cette architecture est plus sure qu'un montage direct sans segmentation.",
        fill=LIGHT_BLUE,
    )
    doc.add_page_break()


def intro(doc):
    doc.add_heading("1. Idee generale de DECEPTR", level=1)
    para(
        doc,
        "DECEPTR est une plateforme de cyberdeception. Le principe est de creer des ressources qui paraissent interessantes pour un attaquant: serveurs, ports, faux portail intranet, fichiers leurres et honeytokens. Quand une personne ou un bot interagit avec ces elements, le systeme collecte les traces et les transforme en alertes utiles pour le SOC.",
    )
    callout(
        doc,
        "Objectif principal",
        "Detecter plus tot les comportements suspects, comprendre les techniques de l'attaquant, conserver les preuves et donner une vue claire au SOC sans exposer les vrais systemes critiques.",
        fill=LIGHT_GREEN,
    )
    doc.add_heading("1.1 Les trois vues de l'architecture", level=2)
    table(
        doc,
        ["Vue", "Question a laquelle elle repond", "Ce qu'elle montre"],
        [
            ["Architecture reseau", "Ou placer chaque composant ?", "Zones, DMZ, LAN analyse, SOC, pare-feu, ports et flux."],
            ["Architecture fonctionnelle", "Comment circule la donnee ?", "Capture, collecte, normalisation, enrichissement, scoring, alertes et dashboards."],
            ["Architecture technique", "Quels services tournent vraiment ?", "Conteneurs Docker, ports, stockage, API, dashboard et validation E2E."],
        ],
        widths=[1.7, 2.4, 2.8],
    )


def network_arch(doc):
    doc.add_heading("2. Architecture reseau", level=1)
    para(doc, "L'architecture reseau explique comment separer les attaquants, les honeypots, les serveurs d'analyse et les postes SOC. C'est la vue la plus importante pour prouver que le honeypot ne met pas en danger le systeme reel.")
    doc.add_heading("2.1 Etapes reseau", level=2)
    table(
        doc,
        ["Etape", "Element", "Objectif", "Explication simple"],
        [
            ["1", "Zone externe", "Representer Internet et les attaquants", "Les bots, scanners et hackers arrivent depuis l'exterieur. Ils ne doivent jamais aller directement vers le LAN interne."],
            ["2", "Perimetre pfSense/FortiGate", "Filtrer et controler l'entree", "Le pare-feu applique NAT, filtrage, IDS/IPS et VPN. Il decide quels ports peuvent atteindre les honeypots."],
            ["3", "DMZ 10.10.10.0/24", "Isoler les leurres exposes", "Cowrie, Web Honeypot et Dionaea sont places dans une zone publique controlee. Si un leurre est attaque, le LAN reste separe."],
            ["4", "T-Pot/Cowrie", "Capturer SSH/Telnet", "Cowrie simule un serveur Linux. Il enregistre usernames, passwords, commandes et sessions."],
            ["5", "Web Honeypot", "Pieger l'acces applicatif", "Le faux portail intranet attire les essais de login et l'acces aux faux documents."],
            ["6", "Dionaea", "Detecter scans et malware", "Dionaea expose FTP, HTTP, SMB et MySQL decoy pour capter les attaques automatisees."],
            ["7", "Switch Core L3", "Relier DMZ et LAN analyse sous controle", "Il symbolise le routage interne. Les flux autorises restent limites."],
            ["8", "LAN Analyse 10.10.20.0/24", "Traiter et stocker les logs", "Logstash, Redis, pipeline, Elasticsearch, MySQL, MinIO, API et dashboard sont ici."],
            ["9", "Admin/SOC 10.10.30.0/24", "Permettre l'exploitation securisee", "Les analystes passent par VPN/MFA et bastion pour consulter Dashboard, Kibana, API et MinIO."],
        ],
        widths=[0.55, 1.45, 1.85, 3.05],
    )
    doc.add_heading("2.2 Pourquoi cette architecture reseau ?", level=2)
    bullets(
        doc,
        [
            "La DMZ protege le LAN: les honeypots sont exposes mais les bases et dashboards restent internes.",
            "Les flux sont unidirectionnels autant que possible: les logs vont des honeypots vers l'analyse, pas l'inverse.",
            "Le SOC accede aux outils via VPN/MFA, ce qui evite d'exposer Kibana, API ou MinIO sur Internet.",
            "Les ports sensibles comme MySQL, Redis et Elasticsearch sont internes ou limites au lab.",
        ],
    )
    doc.add_heading("2.3 Pourquoi pas une architecture plus simple ?", level=2)
    table(
        doc,
        ["Option non retenue", "Pourquoi ce n'est pas ideal"],
        [
            ["Tout mettre sur une seule machine exposee", "Si l'attaquant trouve une faille, il peut atteindre logs, base, dashboard et secrets."],
            ["Exposer Kibana/API directement sur Internet", "Risque de fuite de donnees, brute force et exploitation d'interface d'administration."],
            ["Mettre les honeypots dans le LAN interne", "Un honeypot doit etre considere comme une zone a risque; il ne doit pas partager le meme niveau de confiance que les vrais serveurs."],
            ["Ne pas utiliser VPN/MFA pour le SOC", "Un dashboard cyber contient des informations sensibles et doit rester accessible uniquement aux personnes autorisees."],
        ],
        widths=[2.4, 4.5],
        header_fill=LIGHT_ORANGE,
    )


def functional_arch(doc):
    doc.add_heading("3. Architecture fonctionnelle", level=1)
    para(doc, "L'architecture fonctionnelle explique le chemin de la donnee. Elle commence avec l'action de l'attaquant et se termine avec une alerte exploitable par le SOC.")
    doc.add_heading("3.1 Etapes de traitement", level=2)
    table(
        doc,
        ["Etape", "Module", "Objectif", "Explication simple"],
        [
            ["0", "Filebeat", "Collecter les logs", "Filebeat lit les logs Cowrie et les envoie vers Logstash."],
            ["1", "Logstash TLS 1.3", "Recevoir et parser", "Logstash recoit les evenements de maniere securisee, extrait les champs importants et route les donnees."],
            ["2", "Redis Queue", "Tamponner", "Redis evite la perte d'evenements si beaucoup de logs arrivent en meme temps."],
            ["3", "Collector Python", "Recuperer les evenements", "Le collector lit Redis et prepare les donnees pour le pipeline."],
            ["4", "Normalizer", "Standardiser", "Chaque source est convertie vers un schema commun: IP source, port, username, commande, type d'evenement."],
            ["5", "Enricher", "Ajouter du contexte", "Le systeme ajoute GeoLite2, reputation IP, VirusTotal, AbuseIPDB et Feodo si disponible."],
            ["6", "Correlator", "Relier les evenements", "Plusieurs actions de la meme IP peuvent former une campagne ou une attaque plus claire."],
            ["7", "Risk Scorer", "Calculer le risque", "Le score 0-100 aide a classer LOW, MEDIUM, HIGH ou CRITICAL."],
            ["8", "MITRE Mapper", "Associer ATT&CK", "Exemple: T1110 pour brute force, T1039 pour honeytoken/data access."],
            ["9", "Alert Generator", "Creer une alerte", "Le systeme ecrit une alerte dans MySQL/Elasticsearch et peut notifier ElastAlert."],
        ],
        widths=[0.55, 1.45, 1.7, 3.2],
    )
    doc.add_heading("3.2 Stockage et visualisation", level=2)
    table(
        doc,
        ["Composant", "Ce qu'il garde", "Pourquoi"],
        [
            ["Elasticsearch", "Logs bruts Cowrie, evenements enrichis, web honeypot", "Recherche rapide et visualisation Kibana."],
            ["MySQL", "alerts, iocs, attackers, campaigns, users, api_keys, honeytokens", "Donnees metier structurees pour l'API et le dashboard."],
            ["MinIO", "reports, backups, malware-samples, attachments", "Stockage d'objets et preuves liees aux incidents."],
            ["Kibana", "Data views et explorations", "Analyse avancee des logs et investigation."],
            ["Dashboard DECEPTR", "KPIs, alertes, IoC, campagnes", "Vue simple pour SOC, admin, auditeur et management."],
            ["ElastAlert 2", "Regles et notifications", "Prevenir rapidement par Email/Webhook/Syslog si configure."],
        ],
        widths=[1.7, 2.7, 2.5],
    )
    doc.add_heading("3.3 Pourquoi ce pipeline et pas une lecture directe ?", level=2)
    table(
        doc,
        ["Choix", "Raison"],
        [
            ["Filebeat + Logstash avant le pipeline", "C'est robuste, standard ELK, et adapte aux logs en continu."],
            ["Redis entre Logstash et Python", "Permet d'absorber les pics et d'eviter que le pipeline bloque l'ingestion."],
            ["Pipeline Python separe", "Plus simple pour coder enrichissement, scoring, MITRE et logique metier."],
            ["Elasticsearch + MySQL", "Elasticsearch est meilleur pour chercher dans les logs; MySQL est meilleur pour les donnees metier structurees."],
            ["MinIO separe", "Les fichiers, rapports et preuves ne doivent pas etre melanges avec les logs et tables SQL."],
        ],
        widths=[2.4, 4.5],
        header_fill=LIGHT_GREEN,
    )


def technical_arch(doc):
    doc.add_heading("4. Architecture technique", level=1)
    para(doc, "L'architecture technique explique les conteneurs et les ports reels. Elle sert a montrer que le projet n'est pas seulement theorique: chaque composant peut etre lance, teste et verifie.")
    doc.add_heading("4.1 Conteneurs principaux", level=2)
    table(
        doc,
        ["Conteneur", "Port", "Role simple"],
        [
            ["cowrie", "22 / 23", "Honeypot SSH/Telnet."],
            ["deceptr-web-honeypot", "8082", "Faux portail web intranet."],
            ["deceptr-dionaea", "2121 / 8081 / 1445 / 13306", "Decoys FTP, HTTP, SMB et MySQL."],
            ["deceptr-tpot-forwarder", "sortie vers 5044", "Filebeat qui envoie Cowrie vers Logstash."],
            ["deceptr-logstash", "5044 TLS", "Ingestion logs securisee."],
            ["deceptr-redis", "6379 interne", "Queue d'evenements."],
            ["deceptr-pipeline", "interne", "Analyse Python."],
            ["deceptr-elasticsearch", "9200", "Stockage et recherche logs."],
            ["deceptr-mysql", "3306 interne", "Tables metier."],
            ["deceptr-minio", "9000 / 9001", "Stockage objet et console."],
            ["deceptr-kibana", "5601 HTTP", "Visualisation ELK."],
            ["deceptr-api", "8000 HTTP", "API REST JWT."],
            ["deceptr-dashboard", "8088 HTTP", "Interface SOC."],
            ["deceptr-elastalert", "interne", "Moteur de regles et alertes."],
        ],
        widths=[2.2, 1.55, 3.15],
    )
    doc.add_heading("4.2 Pourquoi Docker ?", level=2)
    bullets(
        doc,
        [
            "Docker rend le projet reproductible: le meme dossier peut etre lance sur une autre machine avec les memes services.",
            "Chaque outil garde ses dependances dans son conteneur, ce qui evite d'installer manuellement Elasticsearch, MySQL, Redis et Logstash dans Windows.",
            "Les volumes Docker gardent les donnees meme si un conteneur redemarre.",
            "Docker Compose permet de decrire l'architecture complete comme du code.",
        ],
    )
    doc.add_heading("4.3 Limites de Docker et partie physique restante", level=2)
    table(
        doc,
        ["Element", "En lab Docker", "En production reelle"],
        [
            ["DMZ", "Simulee par reseaux Docker et ports exposes", "Vraie DMZ avec VLAN, firewall et routage controle."],
            ["pfSense/FortiGate", "Theorique dans le schema", "Equipement ou VM firewall avec regles reelles."],
            ["VPN/MFA", "Documente et recommande", "Obligatoire pour acces SOC/Admin."],
            ["SMTP/SIEM", "Configurable", "A connecter a l'infrastructure officielle."],
            ["Secrets", ".env en lab", "Gestionnaire de secrets ou coffre-fort en production."],
        ],
        widths=[1.5, 2.55, 2.85],
        header_fill=LIGHT_PURPLE,
    )


def conclusion(doc):
    doc.add_heading("5. Conclusion simple", level=1)
    callout(
        doc,
        "Resume",
        "L'architecture reseau protege le systeme par la segmentation. L'architecture fonctionnelle explique comment une action attaquant devient une alerte. L'architecture technique montre les services Docker reels, les ports et la validation. Les trois vues sont necessaires ensemble pour prouver que DECEPTR est securise, comprehensible et deployable.",
        fill=LIGHT_BLUE,
    )
    doc.add_heading("5.1 Phrase courte pour la presentation", level=2)
    para(
        doc,
        "DECEPTR place des leurres dans une DMZ, collecte les actions des attaquants, enrichit les logs avec du renseignement, calcule un score de risque et presente les alertes au SOC sans exposer les vrais systemes internes.",
        bold=True,
        color=NAVY,
    )


def build():
    doc = Document()
    configure(doc)
    cover(doc)
    intro(doc)
    network_arch(doc)
    functional_arch(doc)
    technical_arch(doc)
    conclusion(doc)
    doc.core_properties.title = "DECEPTR - Explication des architectures"
    doc.core_properties.subject = "Architecture reseau, fonctionnelle et technique"
    doc.core_properties.author = "DECEPTR Project"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
